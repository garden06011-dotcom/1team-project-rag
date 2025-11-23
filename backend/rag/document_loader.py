"""
문서 로더 및 청크 분할 모듈

다양한 형식의 문서를 읽고 RAG에 적합한 크기로 분할합니다.
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import os


class Document:
    """문서 데이터 클래스"""

    def __init__(
        self,
        page_content: str,
        metadata: Optional[Dict[str, Any]] = None
    ):
        """
        Args:
            page_content: 문서 내용
            metadata: 메타데이터 (파일명, 페이지 번호 등)
        """
        self.page_content = page_content
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document(content='{self.page_content[:50]}...', metadata={self.metadata})"


class TextLoader:
    """텍스트 파일 로더"""

    def __init__(self, file_path: str, encoding: str = "utf-8"):
        """
        Args:
            file_path: 파일 경로
            encoding: 파일 인코딩
        """
        self.file_path = file_path
        self.encoding = encoding

    def load(self) -> List[Document]:
        """텍스트 파일 로드"""
        try:
            with open(self.file_path, "r", encoding=self.encoding) as f:
                content = f.read()

            metadata = {
                "source": os.path.basename(self.file_path),
                "file_path": self.file_path,
                "file_type": "txt"
            }

            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            print(f"[ERROR] File load failed: {self.file_path}, Error: {e}")
            raise


class PDFLoader:
    """PDF 파일 로더"""

    def __init__(self, file_path: str):
        """
        Args:
            file_path: PDF 파일 경로
        """
        self.file_path = file_path

    def load(self) -> List[Document]:
        """PDF 파일 로드"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(self.file_path)
            documents = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()

                if text.strip():  # 빈 페이지 제외
                    metadata = {
                        "source": os.path.basename(self.file_path),
                        "file_path": self.file_path,
                        "file_type": "pdf",
                        "page": page_num + 1,
                        "total_pages": len(reader.pages)
                    }
                    documents.append(Document(page_content=text, metadata=metadata))

            return documents
        except Exception as e:
            print(f"[ERROR] PDF load failed: {self.file_path}, Error: {e}")
            raise


class DOCXLoader:
    """DOCX 파일 로더"""

    def __init__(self, file_path: str):
        """
        Args:
            file_path: DOCX 파일 경로
        """
        self.file_path = file_path

    def load(self) -> List[Document]:
        """DOCX 파일 로드"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(self.file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)

            metadata = {
                "source": os.path.basename(self.file_path),
                "file_path": self.file_path,
                "file_type": "docx"
            }

            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            print(f"[ERROR] DOCX load failed: {self.file_path}, Error: {e}")
            raise


class DirectoryLoader:
    """디렉토리 내 모든 문서 로드"""

    def __init__(
        self,
        directory_path: str,
        glob_pattern: str = "**/*",
        supported_extensions: List[str] = None
    ):
        """
        Args:
            directory_path: 디렉토리 경로
            glob_pattern: 파일 검색 패턴
            supported_extensions: 지원하는 확장자 리스트
        """
        self.directory_path = Path(directory_path)
        self.glob_pattern = glob_pattern

        if supported_extensions is None:
            self.supported_extensions = [".txt", ".pdf", ".docx", ".md"]
        else:
            self.supported_extensions = supported_extensions

    def load(self) -> List[Document]:
        """디렉토리 내 모든 문서 로드"""
        documents = []

        for file_path in self.directory_path.glob(self.glob_pattern):
            if file_path.is_file() and file_path.suffix in self.supported_extensions:
                print(f"[LOAD] Loading: {file_path.name}")
                try:
                    loader = self._get_loader(str(file_path))
                    docs = loader.load()
                    documents.extend(docs)
                except Exception as e:
                    print(f"[WARN] File load failed (skipped): {file_path.name}, Error: {e}")
                    continue

        print(f"[OK] Total {len(documents)} documents loaded")
        return documents

    def _get_loader(self, file_path: str):
        """파일 확장자에 따라 적절한 로더 반환"""
        ext = Path(file_path).suffix.lower()

        if ext == ".txt" or ext == ".md":
            return TextLoader(file_path)
        elif ext == ".pdf":
            return PDFLoader(file_path)
        elif ext == ".docx":
            return DOCXLoader(file_path)
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {ext}")


class TextSplitter:
    """텍스트 청크 분할기"""

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        separator: str = "\n\n"
    ):
        """
        Args:
            chunk_size: 청크 크기 (문자 수)
            chunk_overlap: 청크 간 오버랩 크기
            separator: 구분자 (문단, 문장 등)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """문서 리스트를 청크로 분할"""
        split_docs = []

        for doc in documents:
            chunks = self._split_text(doc.page_content)

            for i, chunk in enumerate(chunks):
                # 메타데이터 복사 및 청크 정보 추가
                chunk_metadata = doc.metadata.copy()
                chunk_metadata["chunk_index"] = i
                chunk_metadata["total_chunks"] = len(chunks)

                split_docs.append(
                    Document(page_content=chunk, metadata=chunk_metadata)
                )

        return split_docs

    def _split_text(self, text: str) -> List[str]:
        """텍스트를 청크로 분할"""
        if not text or not text.strip():
            return []

        # 구분자로 먼저 분할
        splits = text.split(self.separator)
        chunks = []
        current_chunk = ""

        for split in splits:
            # 현재 청크에 추가했을 때 크기 확인
            if len(current_chunk) + len(split) + len(self.separator) <= self.chunk_size:
                current_chunk += split + self.separator
            else:
                # 현재 청크 저장
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())

                # 새 청크 시작 (오버랩 고려)
                if self.chunk_overlap > 0 and current_chunk:
                    # 이전 청크의 끝부분을 포함
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + split + self.separator
                else:
                    current_chunk = split + self.separator

        # 마지막 청크 저장
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # 너무 긴 청크는 강제로 분할
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > self.chunk_size * 1.5:  # 1.5배 초과 시
                # 문자 단위로 강제 분할
                for i in range(0, len(chunk), self.chunk_size):
                    final_chunks.append(chunk[i:i + self.chunk_size])
            else:
                final_chunks.append(chunk)

        return final_chunks

